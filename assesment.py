import pdfplumber
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ------------------ Helper functions ------------------

def is_bold(word):
    return "Bold" in word.get("fontname", "")

def is_centered(line, page_center):
    x0, x1 = line["x0"], line["x1"]
    return abs((x0 + x1)/2 - page_center) < 20

def get_words_for_line(line, words, threshold=3):
    """Return all words belonging to this line (by top coordinate)"""
    return [w for w in words if abs(w["top"] - line["top"]) < threshold]

def is_line_inside_table(line, table_bbox, tol=2):
    """Check if a line belongs to a table"""
    x0, top, x1, bottom = table_bbox
    return line["top"] >= top - tol and line["bottom"] <= bottom + tol

# ------------------ PDF extraction ------------------

def extract_words_and_lines(pdf_path):
    with pdfplumber.open(pdf_path) as pdf:
        page = pdf.pages[0]
        words = page.extract_words(extra_attrs=["fontname", "size"])
        lines = page.extract_text_lines()
        tables = page.extract_tables()
        table_bboxes = [tb.bbox for tb in page.find_tables()]
    return words, lines, page, tables, table_bboxes

# ------------------ Word creation ------------------

def create_word():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.5)
    section.left_margin = section.right_margin = Inches(0.5)
    return doc

def add_line_paragraph(doc, line, line_words, page_center, font_size=11, font_name="Times New Roman"):
    """Add a paragraph for one line"""
    p = doc.add_paragraph()
    # Alignment
    if is_centered(line, page_center):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Add words
    for word in line_words:
        run = p.add_run(word["text"] + " ")
        run.font.bold = is_bold(word)
        run.font.size = Pt(font_size)
        run.font.name = font_name
    return p

# ------------------ Table handling ------------------

def add_table_to_doc(doc, table, table_words, col_widths=[1, 3, 8]):
    """Add a table to Word doc with merged empty cells, bold/center, and manual column widths"""
    if not table:
        return

    rows = len(table)
    cols = len(table[0])
    word_table = doc.add_table(rows=rows, cols=cols)
    word_table.style = "Table Grid"

    for r in range(rows):
        merge_count = 0
        for c in range(cols):
            value = table[r][c] if table[r][c] else None
            cell = word_table.rows[r].cells[c]

            # Set column width
            cell.width = Cm(col_widths[c])

            if value is None:
                if c > 0:
                    prev_cell = word_table.rows[r].cells[c - 1 - merge_count]
                    prev_cell.merge(cell)
                    merge_count += 1
            else:
                # Set text
                cell.text = value

                # Apply bold / center from table_words
                for w in table_words:
                    if w["text"].strip() == value.strip():
                        for paragraph in cell.paragraphs:
                            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                            run.bold = is_bold(w)
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        break
                merge_count = 0  # reset

    doc.add_paragraph()  # spacing after table

# ------------------ Main ------------------

def main(pdf_path="django_assignment.pdf", output_doc="word_doc.docx"):
    words, lines, page, tables, table_bboxes = extract_words_and_lines(pdf_path)
    doc = create_word()
    page_center = page.width / 2

    # ------------------ Add normal text ------------------
    for line in lines:
        # Skip lines inside tables
        inside_table = any(is_line_inside_table(line, tb_bbox) for tb_bbox in table_bboxes)
        if inside_table:
            continue
        line_words = get_words_for_line(line, words)
        if line_words:
            add_line_paragraph(doc, line, line_words, page_center)

     # ------------------ Add tables first ------------------
    for table in tables:
        # Get words that belong to this table
        table_words = []
        for w in words:
            for tb_bbox in table_bboxes:
                if is_line_inside_table({"top": w["top"], "bottom": w["bottom"]}, tb_bbox):
                    table_words.append(w)
        add_table_to_doc(doc, table, table_words, col_widths=[1,3,15])

    # Save Word file
    doc.save("word_file3.docx")
    print(f"Word document saved: {output_doc}")

# ------------------ Run ------------------
if __name__ == "__main__":
    main()
